from fastapi import APIRouter, UploadFile, File, HTTPException, Query, Form
from fastapi.responses import FileResponse
from typing import List, Optional
import base64
import os
import re
from datetime import datetime, timezone, timedelta

# Beijing timezone (UTC+8)
_BEIJING_TZ = timezone(timedelta(hours=8))

def _beijing_now() -> datetime:
    return datetime.now(_BEIJING_TZ)

from app.services.project_service import ProjectService
from app.services.parser_service import ParserService
from app.services.joern_service import JoernService
from app.models.project import UploadResponse, ProjectStructure, FileStructure, ExtractTestTargetsResponse, TestTargetStats

router = APIRouter(prefix="/api/project", tags=["project"])

@router.post("/upload", response_model=UploadResponse)
async def upload_project(
    file: UploadFile = File(...),
    project_name: Optional[str] = Form(None)
):
    project_id, name, count = await ProjectService.create_project(file, project_name)
    meta = ProjectService.get_project_meta(project_id)
    return UploadResponse(
        project_id=project_id,
        project_name=name,
        file_count=count,
        status="uploaded",
        source="local",
        language=meta.get("language"),
        test_framework=meta.get("test_framework"),
        dependency_manager=meta.get("dependency_manager"),
        env_source=meta.get("env_source", "none"),
    )

@router.get("/list", response_model=List[UploadResponse])
async def list_projects(portal_project_id: Optional[str] = None):
    # portal_project_id: 由 UniPortal iframe 通过 URL query 传入，限定列表为该工程下的 item
    projects = ProjectService.list_projects(portal_project_id=portal_project_id)
    return [UploadResponse(**p) for p in projects]

@router.delete("/{project_id}")
async def delete_project(project_id: str):
    success = ProjectService.delete_project(project_id)
    if not success:
        raise HTTPException(status_code=404, detail="Project not found or could not be deleted")
    return {"status": "success", "project_id": project_id}

@router.get("/{project_id}/has-design-doc")
async def has_design_doc(project_id: str):
    """判断项目是否包含任意设计文档（供前端决定是否执行第二步注释生成）"""
    return {"has_design_doc": ProjectService.has_design_doc(project_id)}

@router.get("/{project_id}/upstream-status")
async def get_upstream_status(project_id: str):
    """检测上游工具（文档审查、需求追溯）的输出数据是否可用"""
    from app.services.upstream_service import UpstreamService
    return UpstreamService.get_upstream_status(project_id)

@router.get("/{project_id}/structure", response_model=ProjectStructure)
async def get_project_structure(project_id: str):
    files = ProjectService.list_files(project_id)
    project_language = ProjectService.get_project_language(project_id)
    project_framework = ProjectService.get_project_framework(project_id)
    structure = []
    
    for path in files:
        # Use path as file_id (base64 encoded to be safe)
        file_id = base64.urlsafe_b64encode(path.encode()).decode()
        ext = os.path.splitext(path)[1].lstrip('.')  # "c", "h", "json", …

        file_language = "python" if ext == "py" else ("c" if ext in {"c", "h"} else None)

        if ext in {'json', 'toml', 'txt'} or path in {'requirements.txt', 'pyproject.toml', 'pytest.ini', 'setup.py'}:
            # JSON design-doc files: no function parsing, just expose for viewing
            structure.append(FileStructure(
                file_id=file_id,
                path=path,
                file_type=ext or 'text',
                language=file_language,
                functions=[]
            ))
            continue

        try:
            content = ProjectService.get_file_content(project_id, path)
            functions = ParserService.parse_functions(content, file_id, file_path=path, language=file_language or project_language)
        except Exception:
            functions = []
            
        structure.append(FileStructure(
            file_id=file_id,
            path=path,
            file_type=ext or 'c',
            language=file_language or project_language,
            functions=functions
        ))
        
    return ProjectStructure(
        project_id=project_id,
        language=project_language,
        test_framework=project_framework,
        files=structure
    )

@router.get("/{project_id}/test-targets", response_model=ExtractTestTargetsResponse)
async def get_test_targets(
    project_id: str,
    include_optional_static: bool = Query(True),
    include_skipped: bool = Query(False),
    header_strategy: str = Query("basename")
):
    result = ParserService.extract_test_targets(project_id, header_strategy=header_strategy)

    optional_static = result.optional_static if include_optional_static else []
    skipped = result.skipped if include_skipped else []

    return ExtractTestTargetsResponse(
        project_id=project_id,
        must_test=result.must_test,
        optional_static=optional_static,
        skipped=skipped,
        stats=TestTargetStats(
            must_test_count=len(result.must_test),
            optional_static_count=len(optional_static),
            skipped_count=len(skipped)
        )
    )

@router.get("/{project_id}/test-summary")
async def get_test_summary(project_id: str):
    """
    Returns aggregated test results for the project from cache.
    Includes function coverage, pass/fail status, and task links.
    """
    from app.services.cache_service import CacheService
    
    # 1. Get all test targets (to know what we should have)
    targets_res = ParserService.extract_test_targets(project_id)
    must_test = targets_res.must_test
    
    summary = []
    
    for func in must_test:
        # 2. Query Cache for each function
        cache_key = func.qualified_name or func.name
        cache_data = CacheService.get_function_data(project_id, func.source_file, cache_key)
        
        # Build summary item
        item = {
            "function_id": func.function_id,
            "name": func.name,
            "signature": func.signature,
            "source_file": func.source_file,
            "language": func.language or ProjectService.get_project_language(project_id),
            "qualified_name": func.qualified_name,
            "task_id": cache_data.get("latest_task_id"),
            "last_run": cache_data.get("last_execution_time"),
            "compile_success": cache_data.get("compile_success"),
            "passed": cache_data.get("passed", 0),
            "failed": cache_data.get("failed", 0),
            "ignored": cache_data.get("ignored", 0),
            "line_coverage": cache_data.get("line_coverage", 0.0),
            "branch_coverage": cache_data.get("branch_coverage", 0.0),
            "status": "unknown"
        }
        
        # Determine Status
        if not item["task_id"]:
            item["status"] = "pending"
        elif item["compile_success"] is False:
            item["status"] = "compile_error"
        elif item["failed"] > 0:
            item["status"] = "failed"
        elif item["passed"] > 0:
            item["status"] = "passed"
        elif item["ignored"] > 0:
            item["status"] = "ignored"
        else:
            item["status"] = "no_tests"
            
        summary.append(item)
        
    # Sort by coverage descending, then by status
    summary.sort(key=lambda x: (x["line_coverage"] or 0, x["passed"] or 0), reverse=True)
    
    return {
        "project_id": project_id,
        "total_functions": len(must_test),
        "functions": summary
    }

@router.get("/{project_id}/export")
async def export_test_results(project_id: str, portal_project_id: str = Query(None)):
    """
    导出项目所有函数信息及测试结果为 JSON（含生成的测试代码）。
    供总览看板的「导出」按钮调用。
    同时写入共享卷 {portal_project_id}/{project_id}/unit-test-generate.json
    """
    from app.services.cache_service import CacheService

    summary = await get_test_summary(project_id)

    # 在汇总结果基础上补充每个函数生成的测试代码、缓存更新时间以及报错详情
    from app.services.runner_service import RunnerService
    for item in summary["functions"]:
        cache_data = CacheService.get_function_data(project_id, item["source_file"], item.get("qualified_name") or item["name"])
        item["test_code"] = cache_data.get("test_code")
        item["updated_at"] = cache_data.get("updated_at")

        # 从 task metadata 读取报错信息
        task_id = item.get("task_id")
        if task_id:
            task_meta = RunnerService.get_task_metadata(task_id)
            if task_meta:
                if task_meta.get("error"):
                    item["compile_error_msg"] = task_meta["error"].get("compile_error", "")
                    item["error_stage"] = task_meta["error"].get("stage", "")
                if task_meta.get("result"):
                    item["test_stdout"] = task_meta["result"].get("stdout", "")
            # 兜底：compile_success 为 False 但 metadata 里没有 error 记录（旧 task）
            if item.get("compile_success") is False and not item.get("compile_error_msg"):
                item["compile_error_msg"] = "编译失败（该任务执行于旧版本，详细错误信息未持久化，请重新执行测试以获取错误详情）"

    funcs = summary["functions"]
    stats = {
        "total": summary["total_functions"],
        "passed": sum(1 for f in funcs if f["status"] == "passed"),
        "failed": sum(1 for f in funcs if f["status"] in ("failed", "ignored")),
        "compile_error": sum(1 for f in funcs if f["status"] == "compile_error"),
    }

    result = {
        "project_id": project_id,
        "project_name": ProjectService.get_project_name(project_id),
        "exported_at": _beijing_now().isoformat(),
        "stats": stats,
        "functions": funcs,
    }

    # 同步写入共享卷（与 project_name 同级，而非嵌套在其内部）:
    # {storage}/{portal_project_id}/{project_id}/unit-test-generate/unit-test-generate.json
    storage = os.getenv("UNIPORTAL_STORAGE_PATH")
    if storage and os.path.isdir(storage) and portal_project_id:
        import json as _json
        item_dir = os.path.join(storage, portal_project_id, project_id)
        out_dir = os.path.join(item_dir, "unit-test-generate")
        os.makedirs(out_dir, exist_ok=True)
        out_path = os.path.join(out_dir, "unit-test-generate.json")
        try:
            with open(out_path, "w", encoding="utf-8") as f:
                _json.dump(result, f, ensure_ascii=False, indent=2)
            print(f"[export] Saved: {out_path}", flush=True)
        except Exception as e:
            print(f"[export] Failed to write shared output: {e}", flush=True)

    return result


# ---- 测试用例拆分 / 注释中文化（供 DOCX 导出使用）----
_TESTCASE_COMMENT_MAP = [
    (re.compile(r'^\s*#\s*objective\s*[:：]\s*', re.IGNORECASE), '测试目标：'),
    (re.compile(r'^\s*#\s*preconditions?\s*[:：]\s*', re.IGNORECASE), '前置条件：'),
    (re.compile(r'^\s*#\s*expected\s*results?\s*[:：]\s*', re.IGNORECASE), '预期结果：'),
    (re.compile(r'^\s*#\s*steps?\s*[:：]\s*', re.IGNORECASE), '测试步骤：'),
    (re.compile(r'^\s*#\s*note\s*[:：]\s*', re.IGNORECASE), '备注：'),
]


def _ensure_period(s: str) -> str:
    s = s.rstrip()
    if s and s[-1] not in '。.；;！!？?':
        s += '。'
    return s


def _comments_to_chinese(comment_lines) -> str:
    """把测试函数前的英文注释（Objective/Preconditions/Expected Results…）
    转成一段可读的中文描述。"""
    parts = []
    for raw in comment_lines:
        for pat, label in _TESTCASE_COMMENT_MAP:
            m = pat.match(raw)
            if m:
                content = raw[m.end():].strip()
                parts.append(label + _ensure_period(content))
                break
        else:
            text = raw.lstrip().lstrip('#').strip()
            if text:
                parts.append(_ensure_period(text))
    return ' '.join(parts)


def _split_test_cases(code: str):
    """把测试源码按顶层 def 拆成独立用例，返回 [{name, comments, body}]。
    每个用例的 comments 为紧贴 def 上方的连续 # 注释行。"""
    if not code:
        return []
    lines = code.splitlines()
    bounds = [i for i, ln in enumerate(lines) if re.match(r'^def\s+\w+\s*\(', ln)]
    if not bounds:
        return []
    cases = []
    for k, start in enumerate(bounds):
        end = bounds[k + 1] if k + 1 < len(bounds) else len(lines)
        m = re.match(r'^def\s+(\w+)\s*\(', lines[start])
        name = m.group(1) if m else f'case_{k + 1}'
        # 向上收集紧贴的注释块（允许中间有空行）
        comments = []
        j = start - 1
        blanks = 0
        while j >= 0:
            stripped = lines[j].strip()
            if stripped == '':
                blanks += 1
                if blanks > 2:
                    break
                j -= 1
                continue
            if stripped.startswith('#'):
                comments.append(lines[j])
                blanks = 0
                j -= 1
                continue
            break
        comments.reverse()
        body_lines = lines[start:end]
        while body_lines and body_lines[-1].strip() == '':
            body_lines.pop()
        cases.append({
            'name': name,
            'comments': comments,
            'body': '\n'.join(body_lines),
        })
    return cases


@router.get("/{project_id}/export-docx")
async def export_test_results_docx(project_id: str):
    """
    导出项目所有函数信息及测试结果为 DOCX 文档（base64 编码）。
    供总览看板的「导出 DOCX」按钮调用。
    """
    from app.services.cache_service import CacheService
    from app.services.llm_service import LLMService
    import io
    import base64 as _b64

    summary = await get_test_summary(project_id)
    language = ProjectService.get_project_language(project_id)

    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx not installed")

    doc = Document()

    # Set default style with CJK font support
    LLMService._set_style_font(doc.styles['Normal'], western="Consolas", east_asian="Microsoft YaHei", size=11)

    # Title
    title = doc.add_heading('项目测试报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        LLMService._set_run_font(run, western="Consolas", east_asian="Microsoft YaHei", size=22)

    LLMService._add_paragraph_with_font(doc, f"项目: {ProjectService.get_project_name(project_id)}", size=11)
    LLMService._add_paragraph_with_font(doc, f"语言: {'Python' if language == 'python' else 'C'}", size=11)
    LLMService._add_paragraph_with_font(doc, f"导出时间: {_beijing_now().strftime('%Y-%m-%d %H:%M:%S')}", size=11)
    doc.add_paragraph()

    funcs = summary.get("functions", [])

    # Helper to set font on all cells in a table
    def _set_table_font(table):
        for row_obj in table.rows:
            for cell in row_obj.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        LLMService._set_run_font(r, western="Consolas", east_asian="Microsoft YaHei", size=9)

    # Stats table
    doc.add_heading('总体统计', level=1)
    passed = sum(1 for f in funcs if f["status"] == "passed")
    failed = sum(1 for f in funcs if f["status"] in ("failed", "ignored"))
    compile_err = sum(1 for f in funcs if f["status"] == "compile_error")
    total = summary.get("total_functions", 1) or 1
    stats_table = doc.add_table(rows=5, cols=2, style='Light Grid Accent 1')
    stats_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (label, val) in enumerate([
        ("总被测函数", str(summary.get("total_functions", 0))),
        ("通过", str(passed)),
        ("失败", str(failed)),
        ("编译错误", str(compile_err)),
        ("通过率", f"{round(passed / total * 100, 1)}%"),
    ]):
        stats_table.rows[i].cells[0].text = label
        stats_table.rows[i].cells[1].text = val
    _set_table_font(stats_table)
    doc.add_paragraph()

    # Per-function details
    doc.add_heading('函数测试详情', level=1)
    for i, func in enumerate(funcs, 1):
        doc.add_heading(f'{i}. {func["name"]}', level=2)

        info_table = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
        status_map = {"passed": "通过", "failed": "失败", "ignored": "失败", "compile_error": "编译错误", "pending": "等待中", "no_tests": "无测试"}
        for j, (label, val) in enumerate([
            ("源文件", func.get("source_file", "")),
            ("签名", func.get("signature", "")),
            ("状态", status_map.get(func.get("status", ""), func.get("status", ""))),
            ("行覆盖率", f"{round((func.get('line_coverage', 0) or 0) * 100, 1)}%"),
            ("分支覆盖率", f"{round((func.get('branch_coverage', 0) or 0) * 100, 1)}%"),
            ("通过/失败/总计", f"{func.get('passed', 0)}/{func.get('failed', 0)}/{func.get('total_tests', 0)}"),
        ]):
            info_table.rows[j].cells[0].text = label
            info_table.rows[j].cells[1].text = val
        _set_table_font(info_table)

        # Test code → 拆分为独立用例，注释转中文段落分别展示
        cache_data = CacheService.get_function_data(project_id, func["source_file"], func.get("qualified_name") or func["name"])
        test_code = cache_data.get("test_code")
        if test_code:
            cases = _split_test_cases(test_code)
            if cases:
                doc.add_heading('测试用例', level=3)
                for ci, case in enumerate(cases, 1):
                    # 用例标题
                    doc.add_heading(f'用例 {ci}：{case["name"]}', level=4)
                    # 注释 → 可读中文段落
                    desc = _comments_to_chinese(case["comments"])
                    if desc:
                        LLMService._add_paragraph_with_font(doc, desc, size=10)
                    # 该用例的代码
                    p = doc.add_paragraph()
                    run = p.add_run(case["body"][:4000])
                    LLMService._set_run_font(run, western="Consolas", east_asian="Microsoft YaHei", size=8)
                    doc.add_paragraph()
            else:
                # 无法按 def 拆分时兜底：整块展示
                doc.add_heading('测试代码', level=3)
                p = doc.add_paragraph()
                run = p.add_run(test_code[:5000])
                LLMService._set_run_font(run, western="Consolas", east_asian="Microsoft YaHei", size=8)

        doc.add_paragraph()

    doc.add_paragraph()
    p_footer = doc.add_paragraph()
    run_footer = p_footer.add_run('文档由单元测试用例智能生成工具自动生成')
    run_footer.italic = True
    LLMService._set_run_font(run_footer, western="Consolas", east_asian="Microsoft YaHei", size=9)

    buf = io.BytesIO()
    doc.save(buf)
    content = _b64.b64encode(buf.getvalue()).decode('ascii')
    safe_name = re.sub(r'[^\w\-.]', '_', ProjectService.get_project_name(project_id))
    filename = f"test_report_{safe_name}_{_beijing_now().strftime('%Y%m%d_%H%M%S')}.docx"

    return {
        "content": content,
        "format": "docx",
        "filename": filename,
    }


@router.get("/{project_id}/file")
async def get_file_source(project_id: str, file_id: str = Query(...)):
    try:
        path = base64.urlsafe_b64decode(file_id).decode()
    except:
        raise HTTPException(status_code=400, detail="Invalid file_id")
        
    content = ProjectService.get_file_content(project_id, path)
    return {
        "file_id": file_id,
        "path": path,
        "content": content
    }

@router.get("/{project_id}/graph/{image_name}")
async def get_project_graph(project_id: str, image_name: str):
    """
    Serves a generated graph image for a project.
    图片生成在本工具私有可写目录中，不在只读的共享卷里。
    """
    local_dir = ProjectService.get_local_project_dir(project_id)
    image_path = os.path.join(local_dir, "graphs", image_name)

    if not os.path.exists(image_path):
        raise HTTPException(status_code=404, detail="Graph image not found")

    return FileResponse(image_path)

@router.get("/{project_id}/function/{function_id}")
async def get_function_detail(project_id: str, function_id: str, use_joern: bool = Query(False), refresh: bool = Query(False)):
    # function_id format: file_id_startline
    try:
        parts = function_id.rsplit('_', 1)
        if len(parts) != 2:
            raise ValueError
        file_id = parts[0]
        start_line = int(parts[1])
        
        path = base64.urlsafe_b64decode(file_id).decode()
    except:
        raise HTTPException(status_code=400, detail="Invalid function_id format")
        
    content = ProjectService.get_file_content(project_id, path)
    functions = ParserService.parse_functions(content, file_id, file_path=path, language=ProjectService.get_project_language(project_id))
    
    target_func = None
    for f in functions:
        if f.start_line == start_line:
            target_func = f
            break
            
    if not target_func:
        raise HTTPException(status_code=404, detail="Function not found")
        
    # Extract source code for function
    lines = content.split('\n')
    # lines are 0-indexed, start_line is 1-indexed
    func_lines = lines[target_func.start_line-1 : target_func.end_line]
    source_code = "\n".join(func_lines)
    
    language = ProjectService.get_project_language(project_id)
    if use_joern and language == "c":
        code_graph = await ParserService.generate_joern_graph(project_id, target_func.name, depth=3, refresh=refresh)
        # Fallback to regex if Joern fails or returns empty
        if not code_graph or not code_graph.get("calls"):
            code_graph = ParserService.generate_code_graph(source_code, language=language)
    else:
        code_graph = ParserService.generate_code_graph(source_code, language=language)
    
    return {
        "function_id": function_id,
        "name": target_func.name,
        "qualified_name": target_func.qualified_name,
        "signature": target_func.signature,
        "start_line": target_func.start_line,
        "end_line": target_func.end_line,
        "source_code": source_code,
        "language": language,
        "class_name": target_func.class_name,
        "is_method": target_func.is_method,
        "is_async": target_func.is_async,
    }

@router.get("/{project_id}/function/{function_id}/requirement-context")
async def get_function_requirement_context(project_id: str, function_id: str):
    """
    查询上游追溯工具中与该函数关联的需求上下文，供前端展示。
    """
    from app.services.upstream_service import UpstreamService
    try:
        parts = function_id.rsplit('_', 1)
        file_id = parts[0]
        start_line = int(parts[1])
        path = base64.urlsafe_b64decode(file_id).decode()
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid function_id")

    language = ProjectService.get_project_language(project_id)
    content = ProjectService.get_file_content(project_id, path)
    functions = ParserService.parse_functions(content, file_id, file_path=path, language=language)
    target_func = next((f for f in functions if f.start_line == start_line), None)
    if not target_func:
        raise HTTPException(status_code=404, detail="Function not found")

    req_ctx = UpstreamService.get_requirement_context(
        project_id,
        target_func.name,
        path,
        signature=getattr(target_func, "signature", None),
        qualified_name=getattr(target_func, "qualified_name", None),
    )
    return req_ctx or {}

@router.get("/{project_id}/function/{function_id}/graph")
async def get_function_graph(project_id: str, function_id: str, use_joern: bool = Query(False), refresh: bool = Query(False)):
    # Legacy endpoint for backward compatibility or bulk fetch
    try:
        parts = function_id.rsplit('_', 1)
        if len(parts) != 2: raise ValueError
        file_id = parts[0]
        start_line = int(parts[1])
        path = base64.urlsafe_b64decode(file_id).decode()
    except:
        raise HTTPException(status_code=400, detail="Invalid function_id format")
        
    content = ProjectService.get_file_content(project_id, path)
    language = ProjectService.get_project_language(project_id)
    functions = ParserService.parse_functions(content, file_id, file_path=path, language=language)
    target_func = next((f for f in functions if f.start_line == start_line), None)
            
    if not target_func:
        raise HTTPException(status_code=404, detail="Function not found")
        
    lines = content.split('\n')
    func_lines = lines[target_func.start_line-1 : target_func.end_line]
    source_code = "\n".join(func_lines)
    
    if use_joern and language == "c":
        code_graph = await ParserService.generate_joern_graph(project_id, target_func.name, depth=3, refresh=refresh)
        if not code_graph or (not code_graph.get("calls") and not code_graph.get("graph_image")):
            code_graph = ParserService.generate_code_graph(source_code, language=language)
    else:
        code_graph = ParserService.generate_code_graph(source_code, language=language)
    
    return code_graph

@router.get("/{project_id}/function/{function_id}/graph/{graph_type}")
async def get_specific_graph(project_id: str, function_id: str, graph_type: str, refresh: bool = Query(False)):
    """
    Generate and get a specific type of graph: call, ast, cfg, pdg
    """
    try:
        parts = function_id.rsplit('_', 1)
        file_id = parts[0]
        start_line = int(parts[1])
        path = base64.urlsafe_b64decode(file_id).decode()
    except:
        raise HTTPException(status_code=400, detail="Invalid function_id format")

    language = ProjectService.get_project_language(project_id)
    if language != "c":
        raise HTTPException(status_code=400, detail="Graph images are currently supported only for C projects")

    content = ProjectService.get_file_content(project_id, path)
    functions = ParserService.parse_functions(content, file_id, file_path=path, language=language)
    target_func = next((f for f in functions if f.start_line == start_line), None)
    if not target_func:
        raise HTTPException(status_code=404, detail="Function not found")

    # 源码路径（可能只读），图谱/CPG 写到本工具私有可写目录
    source_dir = ProjectService.get_project_path(project_id)
    local_dir = ProjectService.get_local_project_dir(project_id)
    cpg_path = os.path.join(local_dir, "cpg.bin")

    # Joern 解析源码目录，但输出到可写的 local_dir
    await JoernService.parse_project(source_dir, cpg_path)

    image_path = None
    if graph_type == "call":
        image_path = await JoernService.generate_graph_image(local_dir, target_func.name, refresh=refresh)
    elif graph_type == "ast":
        image_path = await JoernService.generate_ast_image(local_dir, target_func.name, refresh=refresh)
    elif graph_type == "cfg":
        image_path = await JoernService.generate_cfg_image(local_dir, target_func.name, refresh=refresh)
    elif graph_type == "pdg":
        image_path = await JoernService.generate_pdg_image(local_dir, target_func.name, refresh=refresh)
    else:
        raise HTTPException(status_code=400, detail="Unsupported graph type")

    # Get mtime for browser caching
    mtime = 0
    if image_path:
        full_path = os.path.join(local_dir, "graphs", image_path)
        if os.path.exists(full_path):
            mtime = int(os.path.getmtime(full_path))

    return {"type": graph_type, "image": image_path, "mtime": mtime}
