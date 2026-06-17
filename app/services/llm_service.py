import os
import re
import json
import asyncio
from typing import Any, AsyncIterator, Dict, List, Optional

from openai import AsyncOpenAI

# OpenAI-compatible API configuration.
#
# Works with any OpenAI-compatible endpoint: OpenAI, DeepSeek, Aliyun Bailian
# (DashScope private deployment), Ollama, vLLM, etc. — just point
# OPENAI_BASE_URL at the desired service.
#
# DEEPSEEK_API_KEY is accepted as a fallback so existing .env files keep
# working without modification.
#
# Held as a module-level dict (not constants) so /api/config/llm can mutate
# the live configuration without restarting the process.
def _read_env_config() -> Dict[str, Any]:
    return {
        "api_key": (
            os.getenv("OPENAI_API_KEY")
            or os.getenv("DEEPSEEK_API_KEY")
            or "sk-placeholder"
        ),
        "base_url": (
            os.getenv("OPENAI_BASE_URL")
            or os.getenv("OPENAI_API_BASE")
            or "https://api.deepseek.com"
        ),
        "model": os.getenv("OPENAI_MODEL", "deepseek-chat"),
        "timeout": float(os.getenv("OPENAI_TIMEOUT", "60")),
    }


_config: Dict[str, Any] = _read_env_config()


def get_config() -> Dict[str, Any]:
    """Return a copy of the current live LLM config."""
    return dict(_config)


def update_config(
    *,
    api_key: Optional[str] = None,
    base_url: Optional[str] = None,
    model: Optional[str] = None,
    timeout: Optional[float] = None,
) -> Dict[str, Any]:
    """
    Update live LLM config in-place. Any field passed as None is left unchanged.
    Returns the resulting config snapshot.
    """
    if api_key is not None:
        _config["api_key"] = api_key
    if base_url is not None:
        _config["base_url"] = base_url
    if model is not None:
        _config["model"] = model
    if timeout is not None:
        _config["timeout"] = float(timeout)
    return get_config()


def _is_mock_mode() -> bool:
    key = _config.get("api_key")
    return not key or key == "sk-placeholder"


# Heuristic detection of "thinking / reasoning" model variants. These models
# emit `delta.reasoning_content` first and only switch to `delta.content` once
# they finish reasoning, so the normal max_tokens budget gets eaten by the
# thought process before any actual answer is produced.
_REASONING_KEYWORDS = (
    "reasoner",
    "reasoning",
    "thinking",
    "flash",      # deepseek-v4-flash, gemini-*-flash-thinking, etc.
    "-r1",
    "-r2",
    "deepseek-r",
    "o1-",
    "o1",
    "o3-",
    "o3",
    "o4-",
    "qwq",        # Qwen QwQ reasoning series
)


def _is_reasoning_model(model: str) -> bool:
    if not model:
        return False
    m = model.lower()
    return any(k in m for k in _REASONING_KEYWORDS)


# Floor max_tokens for reasoning models so the thought process + final answer
# both fit. Configurable via LLM_REASONING_MAX_TOKENS.
REASONING_MAX_TOKENS_FLOOR = int(os.getenv("LLM_REASONING_MAX_TOKENS", "16000"))
REASONING_TIMEOUT_FLOOR = float(os.getenv("LLM_REASONING_TIMEOUT", "300"))


def _adjust_for_reasoning(kwargs: Dict[str, Any]) -> Dict[str, Any]:
    """If the active model is a reasoning model, raise max_tokens / timeout
    floors so the thinking phase doesn't starve the answer phase."""
    model = kwargs.get("model") or _config.get("model", "")
    if not _is_reasoning_model(model):
        return kwargs
    if "max_tokens" in kwargs and kwargs["max_tokens"] is not None:
        original = kwargs["max_tokens"]
        boosted = max(int(original), REASONING_MAX_TOKENS_FLOOR)
        if boosted != original:
            kwargs["max_tokens"] = boosted
            print(
                f"LLM: reasoning model '{model}' detected; "
                f"max_tokens bumped {original} -> {boosted}",
                flush=True,
            )
    if "timeout" in kwargs and kwargs["timeout"] is not None:
        kwargs["timeout"] = max(float(kwargs["timeout"]), REASONING_TIMEOUT_FLOOR)
    else:
        # Apply the floor even when timeout wasn't explicitly passed; the
        # client default (_config["timeout"]) may also be too short.
        if _config["timeout"] < REASONING_TIMEOUT_FLOOR:
            kwargs["timeout"] = REASONING_TIMEOUT_FLOOR
    return kwargs


def _make_client() -> AsyncOpenAI:
    # For reasoning models, also raise the client-level timeout so streaming
    # connections don't get killed mid-thought.
    timeout = _config["timeout"]
    if _is_reasoning_model(_config.get("model", "")):
        timeout = max(timeout, REASONING_TIMEOUT_FLOOR)
    return AsyncOpenAI(
        api_key=_config["api_key"],
        base_url=_config["base_url"],
        timeout=timeout,
    )


class LLMService:
    @staticmethod
    def _assess_complexity(function_code: str) -> tuple:
        """
        Estimate the cyclomatic complexity of a C function by counting
        branching keywords.  Returns (prompt_instruction, max_tokens) tailored
        to the complexity level:
          LOW  (branches 0-2)  → 1 test case,   max_tokens=1200
          MED  (branches 3-6)  → 2-3 test cases, max_tokens=2000
          HIGH (branches 7+)   → 3-5 test cases, max_tokens=3500
        """
        branch_keywords = re.findall(
            r'\b(if|else\s+if|else|for|while|do|switch|case|&&|\|\||break|continue|return)\b',
            function_code
        )
        count = len(branch_keywords)

        if count <= 2:
            level = "LOW"
            max_tokens = 1200
            instruction = (
                "6. Write `setUp` and `tearDown` functions (empty bodies are fine).\n"
                "7. **Complexity: LOW** — The function has very few branches. Write exactly **1 test case** "
                "covering the normal execution path. Keep it short.\n"
                "8. Each test function body should be minimal — 1 `TEST_ASSERT_*` call is enough."
            )
        elif count <= 6:
            level = "MED"
            max_tokens = 2000
            instruction = (
                "6. Write `setUp` and `tearDown` functions (empty bodies are fine).\n"
                "7. **Complexity: MEDIUM** — The function has several branches. Write **2–3 test cases**: "
                "one for the normal path, one or two for important edge/boundary conditions. "
                "Do NOT repeat similar variants.\n"
                "8. Keep each test function body concise — prefer 1–2 `TEST_ASSERT_*` calls per test."
            )
        else:
            level = "HIGH"
            max_tokens = 3500
            instruction = (
                "6. Write `setUp` and `tearDown` functions (empty bodies are fine).\n"
                "7. **Complexity: HIGH** — The function has many branches. Write **3–5 test cases** "
                "that collectively cover: the normal path, key boundary values, and the most important "
                "error/special-case branches. Prioritize paths that differ meaningfully in behavior.\n"
                "8. Test bodies can have multiple `TEST_ASSERT_*` calls when a scenario requires verifying "
                "several side-effects, but avoid redundant assertions."
            )

        if os.getenv("LLM_DEBUG_PROMPT") == "1":
            print(f"DEBUG: function complexity = {level} (branch tokens = {count}, max_tokens = {max_tokens})")
        return instruction, max_tokens

    @staticmethod
    async def _chat_stream(
        messages: List[Dict[str, str]],
        *,
        max_tokens: Optional[int] = None,
        temperature: float = 0,
        top_p: Optional[float] = None,
        timeout: Optional[float] = None,
        tools: Optional[list] = None,
        error_template: str = "/* Error: {e} */",
    ) -> AsyncIterator[str]:
        """
        Stream chat completion content via the OpenAI SDK. Yields incremental
        delta.content strings. Compatible with any OpenAI-style endpoint.
        """
        client = _make_client()
        kwargs: Dict[str, Any] = {
            "model": _config["model"],
            "messages": messages,
            "stream": True,
            "temperature": temperature,
        }
        if max_tokens is not None:
            kwargs["max_tokens"] = max_tokens
        if top_p is not None:
            kwargs["top_p"] = top_p
        if timeout is not None:
            kwargs["timeout"] = timeout
        if tools is not None:
            kwargs["tools"] = tools
        kwargs = _adjust_for_reasoning(kwargs)
        debug_stream = os.getenv("LLM_DEBUG_STREAM") == "1"
        try:
            stream = await client.chat.completions.create(**kwargs)
            chunk_count = 0
            content_count = 0
            reasoning_seen = False
            async for chunk in stream:
                chunk_count += 1
                if not chunk.choices:
                    continue
                delta = chunk.choices[0].delta
                if debug_stream:
                    try:
                        print(f"DEBUG chunk#{chunk_count}: {delta.model_dump()} finish={chunk.choices[0].finish_reason}", flush=True)
                    except Exception:
                        print(f"DEBUG chunk#{chunk_count}: {delta} finish={chunk.choices[0].finish_reason}", flush=True)
                # Some models (e.g. reasoning models) emit their tokens in
                # `reasoning_content` first, then switch to `content` for the
                # final answer. Track whether reasoning was seen so the user
                # gets a clear message if no content arrives.
                if delta and getattr(delta, "reasoning_content", None):
                    reasoning_seen = True
                if delta and delta.content:
                    content_count += 1
                    yield delta.content
            if debug_stream:
                print(f"DEBUG stream done: total_chunks={chunk_count} content_chunks={content_count} reasoning_seen={reasoning_seen}", flush=True)
            if content_count == 0:
                hint = (
                    " (model returned only reasoning_content — likely a reasoning model. "
                    "Increase max_tokens or switch to a non-reasoning model.)"
                    if reasoning_seen
                    else " (model returned no content; check model name / parameters)"
                )
                yield error_template.format(e="empty response" + hint)
        except Exception as e:
            print(f"LLM Stream Call failed: {e}", flush=True)
            yield error_template.format(e=e)

    @staticmethod
    async def _chat_once(
        messages: List[Dict[str, str]],
        *,
        max_tokens: Optional[int] = None,
        temperature: float = 0,
        top_p: Optional[float] = None,
        timeout: Optional[float] = None,
        tools: Optional[list] = None,
    ) -> Optional[str]:
        """
        Non-streaming chat completion. Returns the assistant message content,
        or None on failure.
        """
        client = _make_client()
        kwargs: Dict[str, Any] = {
            "model": _config["model"],
            "messages": messages,
            "stream": False,
            "temperature": temperature,
        }
        if max_tokens is not None:
            kwargs["max_tokens"] = max_tokens
        if top_p is not None:
            kwargs["top_p"] = top_p
        if timeout is not None:
            kwargs["timeout"] = timeout
        if tools is not None:
            kwargs["tools"] = tools
        kwargs = _adjust_for_reasoning(kwargs)
        try:
            response = await client.chat.completions.create(**kwargs)
            return response.choices[0].message.content
        except Exception as e:
            print(f"LLM Call failed: {e}")
            return None

    @staticmethod
    def _build_test_prompt(
        language: str,
        project_context: str,
        function_code: str,
        code_graph: Dict[str, Any],
        file_code: str,
        test_framework: str,
        failure_context: Optional[str],
        function_intent: Optional[str],
        prior_test_code: Optional[str],
        design_doc: Optional[dict],
        requirement_context: Optional[Dict[str, Any]] = None,
    ) -> tuple:
        """Build the (prompt, max_tokens) tuple for test-case generation."""
        if language == "python":
            return LLMService._build_python_test_prompt(
                project_context,
                function_code,
                code_graph,
                file_code,
                test_framework,
                failure_context,
                function_intent,
                prior_test_code,
                design_doc,
                requirement_context,
            )

        failure_info = ""
        if failure_context:
            failure_info = (
                "\n### FAILURE CONTEXT (Fix required):\n"
                "The previous test attempt failed. Please analyze the failure information below "
                "and generate a corrected unit test.\n"
            )
            if prior_test_code:
                failure_info += f"\nPrevious (Failed) Test Code:\n```c\n{prior_test_code}\n```\n"
            failure_info += f"\nFailure Details:\n{failure_context}\n"

        intent_info = ""
        if function_intent:
            intent_info = (
                "\n### FUNCTION INTENT:\n"
                "The user (or system) has provided the following description of what the function "
                "is supposed to do. Use this to guide your test case generation, ensuring you test "
                "the intended behavior and edge cases mentioned.\n"
                f"Intent:\n{function_intent}\n"
            )

        design_doc_info = LLMService._format_design_doc(design_doc) if design_doc else ""
        requirement_info = LLMService._format_requirement_context(requirement_context) if requirement_context else ""
        complexity_instruction, max_tokens = LLMService._assess_complexity(function_code)

        prompt = f"""
You are an expert C developer. Generate a unit test file using the {test_framework} framework for the following C function.

{intent_info}
{design_doc_info}
{requirement_info}
{failure_info}

Target Function:
```c
{function_code}
```

Full File Content (for context, macros, and static dependencies):
```c
{file_code}
```

⚠️ Functions already defined in the source file above (DO NOT redefine any of these in your test file):
{LLMService._extract_function_names_warning(file_code)}

Code Graph Info:
{json.dumps(code_graph, indent=2)}

Project Context (Header/Structure info):
{project_context}

⚠️⚠️⚠️ CONCISENESS IS MANDATORY ⚠️⚠️⚠️
Keep the generated test file as SHORT as possible. Every unnecessary line wastes the token budget and will be TRUNCATED.
- No comments, no explanations, no blank lines between assertions.
- Each test function body: ideally 1-3 lines. One call to the function under test + one TEST_ASSERT_*.
- No helper variables unless strictly required.
- setUp/tearDown MUST have empty bodies: `void setUp(void) {{}}` / `void tearDown(void) {{}}`
- main(): only UNITY_BEGIN(), RUN_TEST() calls, UNITY_END(). Nothing else.

🚨 **MUST INCLUDE `int main(void)`** — the test file is compiled as a standalone executable. Without `main()` it will fail to link. Every `test_*` you write MUST appear as a `RUN_TEST(test_name);` line inside `main()`. Skipping `main()` is NEVER acceptable, even under the conciseness rule.

Requirements:
1. Include "unity.h".
2. **DO NOT** include the source ".c" file (e.g., do not use #include "filename.c"). The backend will handle the linking of source files.
3. If a header file for the target function is provided in the "Project Context", include it. Otherwise, declare the function prototype directly in the test file.
4. **IMPORTANT**: Even if the target function is `static` in the source code, do **NOT** use the `static` keyword when declaring its prototype in the test file. The backend automatically removes `static` from the source to make it testable.
5. **CRITICAL – NO REDEFINITION, NO MOCKS, NO WEAK SYMBOLS**:
   - The **entire source file** is compiled and linked automatically. Every function in it is already available.
   - DO NOT redefine, copy-paste, or stub any function from "Full File Content" — not even with `__attribute__((weak))`. The backend uses `objcopy --weaken-symbol` at the binary level to handle symbol conflicts automatically. Writing `__attribute__((weak))` wrappers in the test file is **unnecessary and wastes token budget**.
   - DO NOT declare `extern` global variables unless they appear in a provided header.
   - The ONLY functions you may define are: `setUp`, `tearDown`, `test_*` cases, and `main()`.
   - To test different behaviors: use different input values, not mock functions.
{complexity_instruction}
9. **MANDATORY `main()`** — Always end the file with:
   ```c
   int main(void) {{
       UNITY_BEGIN();
       RUN_TEST(test_xxx);   // one line per test_* you defined above
       ...
       return UNITY_END();
   }}
   ```
   Files without `main()` will fail to link and are considered invalid output.
10. ⚠️ **SEMANTIC COMMENTS REQUIRED (ISO/IEC/IEEE 29119)**: Before EACH `void test_*` function, write a Chinese comment block using the 29119 three-part structure:
    - // Objective（必需）：测试目标描述 — what this specific test validates
    - // Preconditions（可选）：前置条件 — input values, state setup
    - // Expected Results（可选）：预期结果 — expected return value, side effects
    Format as:
    ```
    // Objective：验证<具体场景>时，<被测行为>
    // Preconditions：<输入参数/前置状态>
    // Expected Results：<预期返回值/副作用>
    void test_xxx(void) {{
        ...
    }}
    ```
    If Preconditions or Expected Results are obvious from the Objective, they can be omitted, but Objective is MANDATORY for every test function. This is mandatory for test documentation.
11. Output ONLY the C code for the test file. Wrap it in a ```c block.
12. Ensure the code compiles with standard Unity setup and includes ALL necessary standard headers:
    - `<stdarg.h>` if you use `va_list`, `va_start`, `va_end`, `va_arg`
    - `<string.h>` if you use `strcmp`, `strlen`, `memcpy` etc.
    - `<stdlib.h>` if you use `malloc`, `free` etc.
    - `<math.h>` or `<float.h>` if the function uses floating-point operations
"""
        if os.getenv("LLM_DEBUG_PROMPT") == "1":
            print(f"DEBUG: LLM Prompt:\n{prompt}")
        return prompt, max_tokens

    @staticmethod
    def _build_python_test_prompt(
        project_context: str,
        function_code: str,
        code_graph: Dict[str, Any],
        file_code: str,
        test_framework: str,
        failure_context: Optional[str],
        function_intent: Optional[str],
        prior_test_code: Optional[str],
        design_doc: Optional[dict] = None,
        requirement_context: Optional[Dict[str, Any]] = None,
    ) -> tuple:
        max_tokens = 2600
        failure_info = ""
        if failure_context:
            failure_info = (
                "\n### FAILURE CONTEXT:\n"
                "The previous pytest attempt failed. Generate a corrected test.\n"
            )
            if prior_test_code:
                failure_info += f"\nPrevious Test Code:\n```python\n{prior_test_code}\n```\n"
            failure_info += f"\nFailure Details:\n{failure_context}\n"

        intent_info = ""
        if function_intent:
            intent_info = f"\n### FUNCTION INTENT:\n{function_intent}\n"

        design_doc_info = LLMService._format_design_doc(design_doc) if design_doc else ""
        requirement_info = LLMService._format_requirement_context(requirement_context) if requirement_context else ""

        prompt = f"""
You are an expert Python test engineer. Generate a {test_framework} test file for the following Python target.

{intent_info}
{design_doc_info}
{requirement_info}
{failure_info}

Target Code:
```python
{function_code}
```

Full File Content:
```python
{file_code}
```

Project Context:
{project_context}

Code Graph Info:
{json.dumps(code_graph, indent=2)}

Requirements:
1. Output ONLY Python code wrapped in a ```python block.
2. Use pytest style tests.
3. Prefer direct imports from the target module.
4. Keep the file concise: 2-4 focused tests are enough.
5. Use pytest.raises for exception cases. If you provide a `match` argument, always use a raw string (r"...") to avoid invalid escape sequence warnings. For literal parentheses in the match pattern, write \\( and \\) inside the raw string.
6. Do not add explanations or markdown outside the code block.
7. Do not redefine the target implementation inside the test file.
8. If the target is a class method, instantiate the class only if necessary.
9. If mocking is required, use unittest.mock.patch sparingly.
10. The test file must be directly executable by pytest.
11. ⚠️ **SEMANTIC COMMENTS REQUIRED (ISO/IEC/IEEE 29119)**: Before EACH `def test_*` function, write a Chinese comment block using the 29119 three-part structure:
    - # Objective（必需）：测试目标描述 — what this specific test validates
    - # Preconditions（可选）：前置条件 — input values, state setup, mock configuration
    - # Expected Results（可选）：预期结果 — expected return value, state change, exception type
    Format as:
    ```
    # Objective：验证<具体场景>时，<被测行为>
    # Preconditions：<输入参数/前置状态/mock设置>
    # Expected Results：<预期返回值/异常类型/副作用>
    def test_xxx():
        ...
    ```
    If Preconditions or Expected Results are obvious from the Objective, they can be omitted, but Objective is MANDATORY for every test function. This is mandatory for test documentation.
"""
        if os.getenv("LLM_DEBUG_PROMPT") == "1":
            print(f"DEBUG: Python LLM Prompt:\n{prompt}")
        return prompt, max_tokens

    _TEST_SYSTEM_MSG = (
        "You are a concise C unit test generator. Produce the shortest possible "
        "valid test file. No comments, no explanations, no extra blank lines. "
        "Minimal code only."
    )

    @staticmethod
    async def generate_test_case_stream(
        language: str,
        project_context: str,
        function_code: str,
        code_graph: Dict[str, Any],
        file_code: str = "",
        test_framework: str = "unity",
        failure_context: str = None,
        function_intent: str = None,
        prior_test_code: str = None,
        design_doc: dict = None,
        requirement_context: dict = None,
    ):
        prompt, max_tokens = LLMService._build_test_prompt(
            language, project_context, function_code, code_graph, file_code, test_framework,
            failure_context, function_intent, prior_test_code, design_doc, requirement_context,
        )

        if _is_mock_mode():
            mock_code = LLMService._generate_mock_test(function_code, language=language)
            for char in mock_code:
                yield char
                await asyncio.sleep(0.001)
            return

        system_msg = (
            "You are a concise Python pytest generator. Produce only valid pytest code."
            if language == "python"
            else LLMService._TEST_SYSTEM_MSG
        )
        messages = [
            {"role": "system", "content": system_msg},
            {"role": "user", "content": prompt},
        ]
        async for chunk in LLMService._chat_stream(
            messages,
            max_tokens=max_tokens,
            temperature=0,
            top_p=0.1,
            error_template="/* Error generating test case: {e} */",
        ):
            yield chunk

    @staticmethod
    async def generate_test_case(
        language: str,
        project_context: str,
        function_code: str,
        code_graph: Dict[str, Any],
        file_code: str = "",
        test_framework: str = "unity",
        failure_context: str = None,
        function_intent: str = None,
        prior_test_code: str = None,
        design_doc: dict = None,
        requirement_context: dict = None,
    ) -> str:
        prompt, max_tokens = LLMService._build_test_prompt(
            language, project_context, function_code, code_graph, file_code, test_framework,
            failure_context, function_intent, prior_test_code, design_doc, requirement_context,
        )

        if _is_mock_mode():
            return LLMService._generate_mock_test(function_code, language=language)

        system_msg = (
            "You are a concise Python pytest generator. Produce only valid pytest code."
            if language == "python"
            else LLMService._TEST_SYSTEM_MSG
        )
        messages = [
            {"role": "system", "content": system_msg},
            {"role": "user", "content": prompt},
        ]
        content = await LLMService._chat_once(
            messages,
            max_tokens=max_tokens,
            temperature=0,
            top_p=0.1,
        )
        if content is None:
            return LLMService._generate_mock_test(function_code, language=language)
        return LLMService._extract_code(content)

    @staticmethod
    def _build_intent_prompt(function_code: str, project_context: str, language: str = "c") -> str:
        if language == "python":
            return f"""
You are an expert Python developer. Analyze the following Python target and describe its validation intention following ISO/IEC/IEEE 29119.

Target Code:
```python
{function_code}
```

Project Context:
{project_context}

Requirements for the output format:
1. Use a three-part structure with the following headers:
   #Objective:
   #Preconditions:
   #Expected Results:
2. Be concise and technical.
3. Output ONLY the formatted intent description text.
"""
        return f"""
You are an expert C developer. Analyze the following C function and describe its validation intention following ISO/IEC/IEEE 29119.

Target Function:
```c
{function_code}
```

Project Context:
{project_context}

Requirements for the output format:
1. Use a three-part structure with the following headers:
   #Objective: (Mandatory) Describe the main goal of the function and what it tests/validates.
   #Preconditions: (Optional) Describe any state or setup required before the function logic starts.
   #Expected Results: (Optional) List the specific outcomes, side effects, or return values expected.
2. Be concise and technical.
3. Output ONLY the formatted intent description text.
"""

    @staticmethod
    async def generate_function_intent_stream(
        function_code: str,
        project_context: str = "",
        language: str = "c",
    ):
        prompt = LLMService._build_intent_prompt(function_code, project_context, language=language)

        if _is_mock_mode():
            mock_intent = "#Objective: Mock intent description for testing purposes.\n#Expected Results: 1. Success."
            for char in mock_intent:
                yield char
                await asyncio.sleep(0.001)
            return

        messages = [
            {"role": "system", "content": "You are a helpful assistant that explains code logic."},
            {"role": "user", "content": prompt},
        ]
        async for chunk in LLMService._chat_stream(
            messages,
            temperature=0.2,
            error_template="Error generating intent: {e}",
        ):
            yield chunk

    @staticmethod
    async def generate_function_intent(
        function_code: str,
        project_context: str = "",
        language: str = "c",
    ) -> str:
        prompt = LLMService._build_intent_prompt(function_code, project_context, language=language)

        if _is_mock_mode():
            return "#Objective: Mock intent description for testing purposes.\n#Expected Results: 1. Success."

        messages = [
            {"role": "system", "content": "You are a helpful assistant that explains code logic."},
            {"role": "user", "content": prompt},
        ]
        content = await LLMService._chat_once(messages, temperature=0.2)
        return content if content is not None else "Error generating intent."

    @staticmethod
    def _extract_code(text: str) -> str:
        match = re.search(r'```c\n(.*?)\n```', text, re.DOTALL)
        if match:
            return match.group(1)
        match = re.search(r'```python\n(.*?)\n```', text, re.DOTALL)
        if match:
            return match.group(1)
        # Try without language specifier
        match = re.search(r'```\n(.*?)\n```', text, re.DOTALL)
        if match:
            return match.group(1)
        return text

    @staticmethod
    def _format_design_doc(design_doc: dict) -> str:
        """
        将 function_design 字典格式化为 LLM prompt 中的结构化文本段落。
        """
        if not design_doc:
            return ""

        lines = ["### Function Design Document (from project specification):"]

        basic = design_doc.get("basic_info", {})
        if basic:
            lines.append(f"- Function Name: {basic.get('function_name', '')}")
            lines.append(f"- Identifier: {basic.get('software_unit_identifier', '')}")
            lines.append(f"- Description: {basic.get('function_desc', '')}")

        io = design_doc.get("io_params", {})
        if io:
            inputs = io.get("input_params", [])
            if inputs:
                lines.append("- Input Parameters:")
                for p in inputs:
                    lines.append(f"    - {p.get('param_name', '')}: {p.get('param_desc', '')}")
            ret_vals = io.get("return_value", [])
            if ret_vals:
                lines.append("- Return Values:")
                for r in ret_vals:
                    lines.append(f"    - {r.get('value', '')}: {r.get('desc', '')}")

        algo = design_doc.get("algorithm", {})
        if algo.get("desc"):
            lines.append(f"- Algorithm: {algo['desc']}")

        flow = design_doc.get("logic_flow", {})
        steps = flow.get("flow_steps", [])
        if steps:
            lines.append("- Logic Flow Steps:")
            for step in steps:
                lines.append(f"    {step}")

        call_rel = design_doc.get("call_relation", {})
        called = call_rel.get("called_functions", [])
        if called:
            lines.append(f"- Called Functions: {', '.join(called)}")

        lines.append("")  # trailing blank line
        return "\n".join(lines)

    @staticmethod
    def _format_requirement_context(req_ctx: dict) -> str:
        """
        将上游需求追溯上下文格式化为 LLM prompt 中的结构化文本段落。
        req_ctx 来自 UpstreamService.get_requirement_context()。
        """
        if not req_ctx:
            return ""

        lines = [
            "### REQUIREMENT CONTEXT (from upstream traceability & document validation tools):",
            f"- Associated Requirement: {req_ctx.get('requirement_label', '')}",
        ]

        if req_ctx.get("requirement_title"):
            lines.append(f"- Requirement Title: {req_ctx['requirement_title']}")

        content = req_ctx.get("requirement_content", "")
        if content:
            lines.append(f"- Requirement Description: {content}")

        tables = req_ctx.get("tables", [])
        if tables:
            lines.append("- Requirement Tables (input/output constraints):")
            for tbl in tables:
                headers = tbl.get("headers", [])
                rows = tbl.get("rows", [])
                if headers:
                    lines.append(f"    Headers: {' | '.join(headers)}")
                for row in rows:
                    lines.append(f"    Row: {' | '.join(row)}")

        if req_ctx.get("similarity") is not None:
            lines.append(f"- Traceability Similarity: {req_ctx['similarity']:.2%}")

        lines.append(
            "\nIMPORTANT: Use the requirement context above to guide test case design. "
            "Ensure tests cover the functional scenarios, input/output constraints, and "
            "boundary conditions described in the requirement."
        )
        lines.append("")  # trailing blank line
        return "\n".join(lines)

    @staticmethod
    def _build_annotate_prompt(test_code: str, design_doc: dict, source_code: str) -> str:
        design_doc_info = LLMService._format_design_doc(design_doc)
        source_section = (
            f"\n被测函数源代码：\n```c\n{source_code}\n```\n" if source_code else ""
        )
        return f"""
你是一位 C 语言测试专家，同时也是代码审查员。下面提供了三份材料：
1. 函数设计文档（规定了函数应有的行为）
2. 被测函数的实际源代码
3. 已生成的 Unity 单元测试代码

{design_doc_info}
{source_section}
已有测试代码：
```c
{test_code}
```

你的任务分为**两个维度**：

【维度A：注释插入】
在每一条 `TEST_ASSERT_*` 断言的紧下方，插入一行中文注释，格式为：
`// [设计文档] 预期：<具体预期值或行为> | 原因：<针对该测试场景的分析——在此输入/条件下，设计文档的哪条规定（算法/流程步骤/返回值说明）决定了该输出>`

注释中"原因"要求：
- 必须结合当前测试场景（测试函数名、传入参数）进行具体分析
- 不能只写"0表示正确"这类泛泛的描述
- 若设计文档未覆盖此场景，写：`原因：设计文档未覆盖此场景，依据代码逻辑推断`

【维度B：缺陷检测】
逐一对比：**设计文档规定的行为** vs **源代码的实际实现**。
- 若发现源代码某处实现与设计文档存在偏差（如条件判断错误、返回值错误、流程顺序错误、边界处理缺失等），则：
  1. 在对应断言的注释末尾追加 ` | ⚠️ 疑似缺陷：<简要描述源代码中哪行/哪个逻辑与设计文档不符>`
  2. 若该缺陷会导致当前断言失败，在注释开头改为 `// [设计文档·缺陷] `
- 若源代码与设计文档完全一致，不需要添加缺陷标注

输出规则：
1. **不要修改任何测试逻辑**，包括断言的参数、测试函数名、setUp/tearDown 等。
2. 输出完整的、带注释的 C 代码文件，用 ```c 代码块包裹。
3. 除了插入注释，**不做任何其他修改**。
"""

    @staticmethod
    def _build_python_annotate_prompt(test_code: str, design_doc: dict, source_code: str) -> str:
        design_doc_info = LLMService._format_design_doc(design_doc)
        source_section = (
            f"\n被测函数源代码：\n```python\n{source_code}\n```\n" if source_code else ""
        )
        return f"""
你是一位 Python 测试专家，同时也是代码审查员。下面提供了三份材料：
1. 函数设计文档（规定了函数应有的行为）
2. 被测函数的实际源代码
3. 已生成的 pytest 单元测试代码

{design_doc_info}
{source_section}
已有测试代码：
```python
{test_code}
```

你的任务分为**两个维度**：

【维度A：注释插入】
在每一条 `assert` 语句的紧下方，插入一行中文注释，格式为：
`# [设计文档] 预期：<具体预期值或行为> | 原因：<针对该测试场景的分析——在此输入/条件下，设计文档的哪条规定（算法/流程步骤/返回值说明）决定了该输出>`

注释中"原因"要求：
- 必须结合当前测试场景（测试函数名、传入参数）进行具体分析
- 不能只写"0表示正确"这类泛泛的描述
- 若设计文档未覆盖此场景，写：`原因：设计文档未覆盖此场景，依据代码逻辑推断`

对于 `pytest.raises(...)` 语句，紧下方注释为：
`# [设计文档] 预期：抛出<异常类型> | 原因：<设计文档中哪条规定决定了该异常>`

【维度B：缺陷检测】
逐一对比：**设计文档规定的行为** vs **源代码的实际实现**。
- 若发现源代码某处实现与设计文档存在偏差（如条件判断错误、返回值错误、流程顺序错误、边界处理缺失等），则：
  1. 在对应 assert 的注释末尾追加 ` | ⚠️ 疑似缺陷：<简要描述源代码中哪行/哪个逻辑与设计文档不符>`
  2. 若该缺陷会导致当前断言失败，在注释开头改为 `# [设计文档·缺陷] `
- 若源代码与设计文档完全一致，不需要添加缺陷标注

输出规则：
1. **不要修改任何测试逻辑**，包括 assert 参数、测试函数名、fixture 等。
2. 输出完整的、带注释的 Python 代码文件，用 ```python 代码块包裹。
3. 除了插入注释，**不做任何其他修改**。
"""

    _ANNOTATE_SYSTEM_MSG = "你是一位专业的 C 语言单元测试工程师。"
    _ANNOTATE_PYTHON_SYSTEM_MSG = "你是一位专业的 Python 单元测试工程师。"
    _ANNOTATE_TIMEOUT = 90.0

    @staticmethod
    def _mock_annotate(test_code: str, language: str = "c") -> str:
        lines = test_code.split('\n')
        result = []
        comment_char = '#' if language == 'python' else '//'
        trigger = 'assert' if language == 'python' else 'TEST_ASSERT'
        for line in lines:
            result.append(line)
            if trigger in line:
                result.append(f'{comment_char} [设计文档] 预期：见设计文档 | 原因：Mock 模式，无 API 密钥')
        return '\n'.join(result)

    @staticmethod
    async def annotate_with_design_doc_stream(
        test_code: str,
        design_doc: dict,
        source_code: str = "",
        language: str = "c"
    ):
        """
        第二步：给已生成的测试代码，基于设计文档在每条断言下插入中文注释。
        同时对比源代码与设计文档，若发现偏差则在注释中标注。
        支持 C (Unity) 和 Python (pytest) 两种语言。
        """
        if language == "python":
            prompt = LLMService._build_python_annotate_prompt(test_code, design_doc, source_code)
            system_msg = LLMService._ANNOTATE_PYTHON_SYSTEM_MSG
            error_template = "# Error annotating test case: {e}"
        else:
            prompt = LLMService._build_annotate_prompt(test_code, design_doc, source_code)
            system_msg = LLMService._ANNOTATE_SYSTEM_MSG
            error_template = "/* Error annotating test case: {e} */"

        if os.getenv("LLM_DEBUG_PROMPT") == "1":
            print(f"DEBUG: Annotate Prompt:\n{prompt}")

        if _is_mock_mode():
            for char in LLMService._mock_annotate(test_code, language=language):
                yield char
                await asyncio.sleep(0.001)
            return

        messages = [
            {"role": "system", "content": system_msg},
            {"role": "user", "content": prompt},
        ]
        async for chunk in LLMService._chat_stream(
            messages,
            temperature=0.2,
            timeout=LLMService._ANNOTATE_TIMEOUT,
            error_template=error_template,
        ):
            yield chunk

    @staticmethod
    async def annotate_with_design_doc(
        test_code: str,
        design_doc: dict,
        source_code: str = "",
        language: str = "c"
    ) -> str:
        """
        第二步（非流式版本）：给已生成的测试代码，基于设计文档在每条断言下插入中文注释。
        同时对比源代码与设计文档，若发现偏差则在注释中标注。
        支持 C (Unity) 和 Python (pytest) 两种语言。
        """
        if language == "python":
            prompt = LLMService._build_python_annotate_prompt(test_code, design_doc, source_code)
            system_msg = LLMService._ANNOTATE_PYTHON_SYSTEM_MSG
        else:
            prompt = LLMService._build_annotate_prompt(test_code, design_doc, source_code)
            system_msg = LLMService._ANNOTATE_SYSTEM_MSG

        if os.getenv("LLM_DEBUG_PROMPT") == "1":
            print(f"DEBUG: Annotate Prompt:\n{prompt}")

        if _is_mock_mode():
            return LLMService._mock_annotate(test_code, language=language)

        messages = [
            {"role": "system", "content": system_msg},
            {"role": "user", "content": prompt},
        ]
        content = await LLMService._chat_once(
            messages,
            temperature=0.2,
            timeout=LLMService._ANNOTATE_TIMEOUT,
        )
        if content is None:
            return test_code  # fallback: return unannotated code
        return LLMService._extract_code(content)

    @staticmethod
    def _extract_function_names_warning(file_code: str) -> str:
        """
        Parse top-level C function definitions from file_code and return a
        formatted warning string listing each name.  Used to tell the LLM
        exactly which symbols it must NOT redefine.
        """
        if not file_code:
            return "(source file is empty)"
        # Match function definitions: identifier followed by '(' at column 0,
        # preceded by a return-type token on the same or previous line.
        # We look for lines that start with an identifier directly (no indent)
        # and are followed by '(' — a reliable heuristic for top-level defs.
        pattern = re.compile(
            r'^(?:(?:static|inline|extern|const|volatile|unsigned|signed|long|short)\s+)*'
            r'(?:struct\s+\w+\s*\*?\s*|enum\s+\w+\s*|union\s+\w+\s*|\w+\s*\*?\s*)'
            r'(?P<fname>[a-zA-Z_]\w*)\s*\(',
            re.MULTILINE
        )
        # Keywords that should never be treated as function names
        KEYWORDS = {
            'if', 'for', 'while', 'switch', 'return', 'sizeof', 'typedef',
            'struct', 'enum', 'union', 'else', 'do', 'goto', 'case', 'default',
            'break', 'continue', 'void', 'int', 'char', 'float', 'double',
            'unsigned', 'signed', 'long', 'short', 'static', 'extern',
            'const', 'volatile', 'inline', 'register', 'auto',
        }
        names = []
        for m in pattern.finditer(file_code):
            fname = m.group('fname')
            if fname not in KEYWORDS and fname not in names:
                names.append(fname)
        if not names:
            return "(no function definitions detected)"
        return "  " + ", ".join(f"`{n}`" for n in names)

    @staticmethod
    def _generate_mock_test(function_code: str, language: str = "c") -> str:
        # Generate a dummy test based on function name
        func_name_match = re.search(r'([a-zA-Z0-9_]+)\s*\(', function_code)
        func_name = func_name_match.group(1) if func_name_match else "unknown"

        if language == "python":
            return f"""
import pytest

def test_{func_name}_basic():
    pytest.skip("Mock test generated without API key")
"""

        return f"""
#include "unity.h"
#include <math.h>
#include <float.h>

// Declare prototype without static (even if static in source)
// The backend handles removing static from the source file.
// double {func_name}(...);

void setUp(void) {{}}
void tearDown(void) {{}}

void test_{func_name}_basic(void) {{
    // TODO: Implement actual test
    TEST_IGNORE_MESSAGE("Mock test generated without API key");
}}

int main(void) {{
    UNITY_BEGIN();
    RUN_TEST(test_{func_name}_basic);
    return UNITY_END();
}}
"""

    # ── Document Export ──────────────────────────────────────────────────

    @staticmethod
    def generate_export_document(
        *,
        language: str,
        function_name: str,
        function_signature: str,
        qualified_name: Optional[str],
        class_name: Optional[str],
        source_file: str,
        source_code: str,
        function_intent: str,
        test_code: str,
        design_doc: Optional[dict],
        test_result: Optional[dict],
        coverage: Optional[dict],
        format: str = "markdown",
    ) -> str:
        """
        Generate a test document (Markdown or HTML) that includes:
        - Function basic info, signature, source code
        - Design document (if available)
        - Function intent (semantic meaning)
        - Test code with annotations
        - Test execution results and coverage
        """
        is_python = language == "python"
        display_name = qualified_name or function_name

        if format == "html":
            return LLMService._build_export_html(
                is_python=is_python,
                display_name=display_name,
                function_name=function_name,
                function_signature=function_signature,
                qualified_name=qualified_name,
                class_name=class_name,
                source_file=source_file,
                source_code=source_code,
                function_intent=function_intent,
                test_code=test_code,
                design_doc=design_doc,
                test_result=test_result,
                coverage=coverage,
            )
        if format == "docx":
            return LLMService._build_export_docx(
                is_python=is_python,
                display_name=display_name,
                function_name=function_name,
                function_signature=function_signature,
                qualified_name=qualified_name,
                class_name=class_name,
                source_file=source_file,
                source_code=source_code,
                function_intent=function_intent,
                test_code=test_code,
                design_doc=design_doc,
                test_result=test_result,
                coverage=coverage,
            )
        return LLMService._build_export_markdown(
            is_python=is_python,
            display_name=display_name,
            function_name=function_name,
            function_signature=function_signature,
            qualified_name=qualified_name,
            class_name=class_name,
            source_file=source_file,
            source_code=source_code,
            function_intent=function_intent,
            test_code=test_code,
            design_doc=design_doc,
            test_result=test_result,
            coverage=coverage,
        )

    @staticmethod
    def _build_export_markdown(
        is_python: bool,
        display_name: str,
        function_name: str,
        function_signature: str,
        qualified_name: Optional[str],
        class_name: Optional[str],
        source_file: str,
        source_code: str,
        function_intent: str,
        test_code: str,
        design_doc: Optional[dict],
        test_result: Optional[dict],
        coverage: Optional[dict],
    ) -> str:
        lang_label = "Python" if is_python else "C"
        lang_tag = "python" if is_python else "c"
        fw_label = "pytest" if is_python else "Unity"

        lines = []
        lines.append(f"# 单元测试用例文档")
        lines.append(f"")
        lines.append(f"## 1. 基本信息")
        lines.append(f"")
        lines.append(f"| 字段 | 内容 |")
        lines.append(f"|------|------|")
        lines.append(f"| 被测函数 | `{display_name}` |")
        if class_name:
            lines.append(f"| 所属类 | `{class_name}` |")
        lines.append(f"| 语言 | {lang_label} |")
        lines.append(f"| 测试框架 | {fw_label} |")
        lines.append(f"| 源文件 | `{source_file}` |")
        lines.append(f"")
        lines.append(f"### 函数签名")
        lines.append(f"")
        lines.append(f"```{lang_tag}")
        lines.append(function_signature)
        lines.append(f"```")
        lines.append(f"")

        # Source code
        if source_code:
            lines.append(f"### 源代码")
            lines.append(f"")
            lines.append(f"```{lang_tag}")
            lines.append(source_code)
            lines.append(f"```")
            lines.append(f"")

        # Design document
        if design_doc:
            lines.append(f"## 2. 设计文档")
            lines.append(f"")
            basic = design_doc.get("basic_info", {})
            if basic:
                lines.append(f"- **函数名称**: {basic.get('function_name', function_name)}")
                lines.append(f"- **软件单元标识**: {basic.get('software_unit_identifier', '-')}")
                desc = basic.get('function_desc', '')
                if desc:
                    lines.append(f"- **功能描述**: {desc}")
                lines.append(f"")
            io_info = design_doc.get("io_params", {})
            inputs = io_info.get("input_params", [])
            if inputs:
                lines.append(f"### 输入参数")
                lines.append(f"")
                for p in inputs:
                    lines.append(f"- **{p.get('param_name', '?')}**: {p.get('param_desc', '-')}")
                lines.append(f"")
            ret_vals = io_info.get("return_value", [])
            if ret_vals:
                lines.append(f"### 返回值")
                lines.append(f"")
                for r in ret_vals:
                    lines.append(f"- **{r.get('value', '?')}**: {r.get('desc', '-')}")
                lines.append(f"")
            algo = design_doc.get("algorithm", {})
            if algo.get("desc"):
                lines.append(f"### 算法描述")
                lines.append(f"")
                lines.append(algo["desc"])
                lines.append(f"")
            flow = design_doc.get("logic_flow", {})
            steps = flow.get("flow_steps", [])
            if steps:
                lines.append(f"### 逻辑流程")
                lines.append(f"")
                for i, step in enumerate(steps, 1):
                    lines.append(f"{i}. {step}")
                lines.append(f"")
            call_rel = design_doc.get("call_relation", {})
            called = call_rel.get("called_functions", [])
            if called:
                lines.append(f"### 调用关系")
                lines.append(f"")
                lines.append(f"被调用函数: {', '.join('`' + c + '`' for c in called)}")
                lines.append(f"")

        # Function intent (semantic meaning)
        lines.append(f"## {3 if design_doc else 2}. 函数意图（语义分析）")
        lines.append(f"")
        if function_intent:
            lines.append(function_intent)
        else:
            lines.append("*(未提供函数意图)*")
        lines.append(f"")

        # Test code
        section_num = 4 if design_doc else 3
        lines.append(f"## {section_num}. 测试用例")
        lines.append(f"")
        lines.append(f"```{lang_tag}")
        lines.append(test_code)
        lines.append(f"```")
        lines.append(f"")

        # Test results
        section_num += 1
        lines.append(f"## {section_num}. 测试执行结果")
        lines.append(f"")
        if test_result:
            passed = test_result.get("passed", 0)
            failed = test_result.get("failed", 0)
            total = test_result.get("total", 0)
            lines.append(f"| 指标 | 数值 |")
            lines.append(f"|------|------|")
            lines.append(f"| 通过 | {passed} |")
            lines.append(f"| 失败 | {failed} |")
            lines.append(f"| 总计 | {total} |")
            if total > 0:
                rate = round(passed / total * 100, 1)
                lines.append(f"| 通过率 | {rate}% |")
            lines.append(f"")
        else:
            lines.append("*(尚未执行测试)*")
            lines.append(f"")

        # Coverage
        if coverage:
            section_num += 1
            lines.append(f"## {section_num}. 覆盖率分析")
            lines.append(f"")
            for fcov in coverage.get("files", []):
                lines.append(f"### {fcov.get('file', 'unknown')}")
                lines.append(f"")
                lines.append(f"| 覆盖类型 | 覆盖行/分支 | 总行/分支 | 覆盖率 |")
                lines.append(f"|----------|-------------|-----------|--------|")
                for key, label in [("line", "行覆盖率"), ("function", "函数覆盖率"), ("branch", "分支覆盖率")]:
                    detail = fcov.get(key)
                    if detail:
                        cov_val = detail.get("covered", 0)
                        tot_val = detail.get("total", 0)
                        r = round(detail.get("rate", 0) * 100, 1)
                        lines.append(f"| {label} | {cov_val} | {tot_val} | {r}% |")
                lines.append(f"")
        else:
            lines.append("")

        lines.append("---")
        lines.append(f"")
        lines.append(f"*文档由单元测试用例智能生成工具自动生成*")
        lines.append(f"")

        return "\n".join(lines)

    @staticmethod
    def _build_export_html(
        is_python: bool,
        display_name: str,
        function_name: str,
        function_signature: str,
        qualified_name: Optional[str],
        class_name: Optional[str],
        source_file: str,
        source_code: str,
        function_intent: str,
        test_code: str,
        design_doc: Optional[dict],
        test_result: Optional[dict],
        coverage: Optional[dict],
    ) -> str:
        """Build an HTML export document (self-contained, minimal styling)."""
        import html as _html
        esc = _html.escape
        lang_label = "Python" if is_python else "C"
        lang_tag = "python" if is_python else "c"
        fw_label = "pytest" if is_python else "Unity"

        parts = []
        parts.append("""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>单元测试用例文档</title>
<style>
body { font-family: -apple-system, 'Segoe UI', system-ui, sans-serif; max-width: 900px; margin: 0 auto; padding: 2em; color: #1e293b; line-height: 1.7; }
h1 { border-bottom: 3px solid #2563eb; padding-bottom: .5em; }
h2 { border-bottom: 2px solid #e2e8f0; padding-bottom: .3em; margin-top: 2em; }
table { border-collapse: collapse; width: 100%; }
th, td { border: 1px solid #e2e8f0; padding: .5em .75em; text-align: left; }
th { background: #f8fafc; }
pre { background: #1e293b; color: #e2e8f0; padding: 1em; border-radius: 8px; overflow-x: auto; }
code { font-family: 'JetBrains Mono', 'Fira Code', monospace; font-size: .9em; }
.pass { color: #16a34a; font-weight: 700; }
.fail { color: #dc2626; font-weight: 700; }
.footer { color: #94a3b8; font-size: .85em; margin-top: 3em; }
</style>
</head>
<body>
""")
        parts.append(f"<h1>单元测试用例文档</h1>")

        # Section 1: Basic info
        parts.append(f"<h2>1. 基本信息</h2>")
        parts.append(f"<table>")
        parts.append(f"<tr><th>被测函数</th><td><code>{esc(display_name)}</code></td></tr>")
        if class_name:
            parts.append(f"<tr><th>所属类</th><td><code>{esc(class_name)}</code></td></tr>")
        parts.append(f"<tr><th>语言</th><td>{lang_label}</td></tr>")
        parts.append(f"<tr><th>测试框架</th><td>{fw_label}</td></tr>")
        parts.append(f"<tr><th>源文件</th><td><code>{esc(source_file)}</code></td></tr>")
        parts.append(f"</table>")

        parts.append(f"<h3>函数签名</h3>")
        parts.append(f"<pre><code>{esc(function_signature)}</code></pre>")

        if source_code:
            parts.append(f"<h3>源代码</h3>")
            parts.append(f"<pre><code>{esc(source_code)}</code></pre>")

        # Section 2: Design doc
        sec = 2
        if design_doc:
            parts.append(f"<h2>{sec}. 设计文档</h2>")
            sec += 1
            basic = design_doc.get("basic_info", {})
            if basic:
                parts.append(f"<p><strong>函数名称:</strong> {esc(basic.get('function_name', function_name))}</p>")
                parts.append(f"<p><strong>软件单元标识:</strong> {esc(basic.get('software_unit_identifier', '-'))}</p>")
                desc = basic.get('function_desc', '')
                if desc:
                    parts.append(f"<p><strong>功能描述:</strong> {esc(desc)}</p>")
            io_info = design_doc.get("io_params", {})
            inputs = io_info.get("input_params", [])
            if inputs:
                parts.append(f"<h3>输入参数</h3><ul>")
                for p in inputs:
                    parts.append(f"<li><strong>{esc(p.get('param_name', '?'))}:</strong> {esc(p.get('param_desc', '-'))}</li>")
                parts.append(f"</ul>")
            ret_vals = io_info.get("return_value", [])
            if ret_vals:
                parts.append(f"<h3>返回值</h3><ul>")
                for r in ret_vals:
                    parts.append(f"<li><strong>{esc(r.get('value', '?'))}:</strong> {esc(r.get('desc', '-'))}</li>")
                parts.append(f"</ul>")

        # Function intent
        parts.append(f"<h2>{sec}. 函数意图（语义分析）</h2>")
        sec += 1
        if function_intent:
            parts.append(f"<pre>{esc(function_intent)}</pre>")
        else:
            parts.append("<p><em>(未提供函数意图)</em></p>")

        # Test code
        parts.append(f"<h2>{sec}. 测试用例</h2>")
        sec += 1
        parts.append(f"<pre><code>{esc(test_code)}</code></pre>")

        # Test results
        parts.append(f"<h2>{sec}. 测试执行结果</h2>")
        sec += 1
        if test_result:
            passed = test_result.get("passed", 0)
            failed = test_result.get("failed", 0)
            total = test_result.get("total", 0)
            parts.append(f"<table>")
            parts.append(f"<tr><th>通过</th><td class='pass'>{passed}</td></tr>")
            parts.append(f"<tr><th>失败</th><td class='fail'>{failed}</td></tr>")
            parts.append(f"<tr><th>总计</th><td>{total}</td></tr>")
            if total > 0:
                rate = round(passed / total * 100, 1)
                parts.append(f"<tr><th>通过率</th><td>{rate}%</td></tr>")
            parts.append(f"</table>")
        else:
            parts.append("<p><em>(尚未执行测试)</em></p>")

        # Coverage
        if coverage:
            parts.append(f"<h2>{sec}. 覆盖率分析</h2>")
            for fcov in coverage.get("files", []):
                parts.append(f"<h3>{esc(fcov.get('file', 'unknown'))}</h3>")
                parts.append(f"<table>")
                parts.append(f"<tr><th>覆盖类型</th><th>覆盖行/分支</th><th>总行/分支</th><th>覆盖率</th></tr>")
                for key, label in [("line", "行覆盖率"), ("function", "函数覆盖率"), ("branch", "分支覆盖率")]:
                    detail = fcov.get(key)
                    if detail:
                        cov_val = detail.get("covered", 0)
                        tot_val = detail.get("total", 0)
                        r = round(detail.get("rate", 0) * 100, 1)
                        parts.append(f"<tr><td>{label}</td><td>{cov_val}</td><td>{tot_val}</td><td>{r}%</td></tr>")
                parts.append(f"</table>")

        parts.append("<hr><p class='footer'>文档由单元测试用例智能生成工具自动生成</p>")
        parts.append("</body></html>")

        return "\n".join(parts)

    @staticmethod
    def _set_run_font(run, western: str = "Consolas", east_asian: str = "Microsoft YaHei", size: int = 11):
        """Set both Western and East Asian fonts on a run for proper CJK rendering."""
        from docx.shared import Pt as _Pt
        from lxml import etree
        rPr = run._r.get_or_add_rPr()
        ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        rFonts = rPr.find(f'{ns}rFonts')
        if rFonts is None:
            rFonts = etree.SubElement(rPr, f'{ns}rFonts')
        rFonts.set(f'{ns}ascii', western)
        rFonts.set(f'{ns}hAnsi', western)
        rFonts.set(f'{ns}eastAsia', east_asian)
        run.font.size = _Pt(size)

    @staticmethod
    def _set_style_font(style, western: str = "Consolas", east_asian: str = "Microsoft YaHei", size: int = 11):
        """Set both Western and East Asian fonts on a paragraph style for proper CJK rendering."""
        from docx.shared import Pt as _Pt
        from lxml import etree
        style.font.name = western
        style.font.size = _Pt(size)
        rPr = style.element.get_or_add_rPr()
        ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        rFonts = rPr.find(f'{ns}rFonts')
        if rFonts is None:
            rFonts = etree.SubElement(rPr, f'{ns}rFonts')
        rFonts.set(f'{ns}ascii', western)
        rFonts.set(f'{ns}hAnsi', western)
        rFonts.set(f'{ns}eastAsia', east_asian)

    @staticmethod
    def _add_paragraph_with_font(doc, text: str, western: str = "Consolas", east_asian: str = "Microsoft YaHei", size: int = 11, style: str = None):
        """Add a paragraph with proper CJK font support."""
        p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
        run = p.add_run(text)
        LLMService._set_run_font(run, western=western, east_asian=east_asian, size=size)
        return p

    @staticmethod
    def _build_export_docx(
        is_python: bool,
        display_name: str,
        function_name: str,
        function_signature: str,
        qualified_name: Optional[str],
        class_name: Optional[str],
        source_file: str,
        source_code: str,
        function_intent: str,
        test_code: str,
        design_doc: Optional[dict],
        test_result: Optional[dict],
        coverage: Optional[dict],
    ) -> str:
        """Build a DOCX test document, returned as base64-encoded bytes."""
        import io
        import base64 as _b64

        try:
            from docx import Document
            from docx.shared import Inches, Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
        except ImportError:
            return ""  # python-docx not installed

        doc = Document()

        # Set default style with CJK font support
        style = doc.styles['Normal']
        LLMService._set_style_font(style, western="Consolas", east_asian="Microsoft YaHei", size=11)

        lang_label = "Python" if is_python else "C"
        fw_label = "pytest" if is_python else "Unity"

        # Title
        title = doc.add_heading('单元测试用例文档', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            LLMService._set_run_font(run, western="Consolas", east_asian="Microsoft YaHei", size=22)

        # Section 1: Basic Info
        h1 = doc.add_heading('1. 基本信息', level=1)
        for run in h1.runs:
            LLMService._set_run_font(run, western="Consolas", east_asian="Microsoft YaHei", size=16)

        table = doc.add_table(rows=5 + (1 if class_name else 0), cols=2, style='Light Grid Accent 1')
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        row = 0
        for label, val in [
            ("被测函数", display_name),
            ("所属类", class_name),
            ("语言", lang_label),
            ("测试框架", fw_label),
            ("源文件", source_file),
        ]:
            if val is None:
                continue
            table.rows[row].cells[0].text = label
            table.rows[row].cells[1].text = str(val)
            for cell in table.rows[row].cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        LLMService._set_run_font(r, western="Consolas", east_asian="Microsoft YaHei", size=10)
            row += 1
        doc.add_paragraph()

        doc.add_heading('函数签名', level=2)
        LLMService._add_paragraph_with_font(doc, function_signature, western="Consolas", east_asian="Microsoft YaHei", size=10, style='No Spacing')

        if source_code:
            doc.add_heading('源代码', level=2)
            LLMService._add_paragraph_with_font(doc, source_code, western="Consolas", east_asian="Microsoft YaHei", size=10, style='No Spacing')

        sec = 2

        # Design doc
        if design_doc:
            doc.add_heading(f'{sec}. 设计文档', level=1)
            sec += 1
            basic = design_doc.get("basic_info", {})
            if basic:
                LLMService._add_paragraph_with_font(doc, f"函数名称: {basic.get('function_name', function_name)}", size=10)
                LLMService._add_paragraph_with_font(doc, f"软件单元标识: {basic.get('software_unit_identifier', '-')}", size=10)
                desc = basic.get('function_desc', '')
                if desc:
                    LLMService._add_paragraph_with_font(doc, f"功能描述: {desc}", size=10)
            io_info = design_doc.get("io_params", {})
            inputs = io_info.get("input_params", [])
            if inputs:
                doc.add_heading('输入参数', level=2)
                for p in inputs:
                    LLMService._add_paragraph_with_font(doc, f"{p.get('param_name', '?')}: {p.get('param_desc', '-')}", size=10)
            ret_vals = io_info.get("return_value", [])
            if ret_vals:
                doc.add_heading('返回值', level=2)
                for r in ret_vals:
                    LLMService._add_paragraph_with_font(doc, f"{r.get('value', '?')}: {r.get('desc', '-')}", size=10)

        # Function intent
        h_intent = doc.add_heading(f'{sec}. 函数意图（语义分析）', level=1)
        for run in h_intent.runs:
            LLMService._set_run_font(run, western="Consolas", east_asian="Microsoft YaHei", size=16)
        sec += 1
        if function_intent:
            LLMService._add_paragraph_with_font(doc, function_intent, size=10)
        else:
            p = doc.add_paragraph()
            run = p.add_run('(未提供函数意图)')
            run.italic = True
            LLMService._set_run_font(run, western="Consolas", east_asian="Microsoft YaHei", size=10)

        # Test code
        doc.add_heading(f'{sec}. 测试用例', level=1)
        sec += 1
        p = doc.add_paragraph()
        run = p.add_run(test_code)
        LLMService._set_run_font(run, western="Consolas", east_asian="Microsoft YaHei", size=9)

        # Test results
        doc.add_heading(f'{sec}. 测试执行结果', level=1)
        sec += 1
        if test_result:
            passed = test_result.get("passed", 0)
            failed = test_result.get("failed", 0)
            total = test_result.get("total", 0)
            res_table = doc.add_table(rows=4, cols=2, style='Light Grid Accent 1')
            for i, (label, val) in enumerate([
                ("通过", str(passed)),
                ("失败", str(failed)),
                ("总计", str(total)),
                ("通过率", f"{round(passed / total * 100, 1)}%" if total > 0 else "N/A"),
            ]):
                res_table.rows[i].cells[0].text = label
                res_table.rows[i].cells[1].text = val
                for cell in res_table.rows[i].cells:
                    for p2 in cell.paragraphs:
                        for r2 in p2.runs:
                            LLMService._set_run_font(r2, western="Consolas", east_asian="Microsoft YaHei", size=10)
        else:
            p = doc.add_paragraph()
            run = p.add_run('(尚未执行测试)')
            run.italic = True
            LLMService._set_run_font(run, western="Consolas", east_asian="Microsoft YaHei", size=10)
        doc.add_paragraph()

        # Coverage
        if coverage:
            doc.add_heading(f'{sec}. 覆盖率分析', level=1)
            for fcov in coverage.get("files", []):
                doc.add_heading(fcov.get('file', 'unknown'), level=2)
                cov_table = doc.add_table(rows=4, cols=4, style='Light Grid Accent 1')
                cov_table.rows[0].cells[0].text = "覆盖类型"
                cov_table.rows[0].cells[1].text = "覆盖行/分支"
                cov_table.rows[0].cells[2].text = "总行/分支"
                cov_table.rows[0].cells[3].text = "覆盖率"
                for i, (key, label) in enumerate([("line", "行覆盖率"), ("function", "函数覆盖率"), ("branch", "分支覆盖率")], 1):
                    detail = fcov.get(key)
                    if detail:
                        cov_table.rows[i].cells[0].text = label
                        cov_table.rows[i].cells[1].text = str(detail.get("covered", 0))
                        cov_table.rows[i].cells[2].text = str(detail.get("total", 0))
                        cov_table.rows[i].cells[3].text = f"{round(detail.get('rate', 0) * 100, 1)}%"
                for row_obj in cov_table.rows:
                    for cell in row_obj.cells:
                        for p2 in cell.paragraphs:
                            for r2 in p2.runs:
                                LLMService._set_run_font(r2, western="Consolas", east_asian="Microsoft YaHei", size=10)

        doc.add_paragraph()
        p_footer = doc.add_paragraph()
        run_footer = p_footer.add_run('文档由单元测试用例智能生成工具自动生成')
        run_footer.italic = True
        LLMService._set_run_font(run_footer, western="Consolas", east_asian="Microsoft YaHei", size=9)

        # Save to bytes and base64 encode
        buf = io.BytesIO()
        doc.save(buf)
        return _b64.b64encode(buf.getvalue()).decode('ascii')
